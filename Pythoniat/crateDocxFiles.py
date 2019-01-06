from docx import Document
from docx.shared import Inches
doc = Document()
doc.add_heading("python programming")
doc.add_paragraph("python is high livel language consider between the best util languages\n to getting started we will write our first script in python \n print('hello world !!')")
doc.save('dem.docx')